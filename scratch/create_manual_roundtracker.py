from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manual"
IMG_DIR = OUT_DIR / "images"
DOCX_PATH = OUT_DIR / "manual_roundtracker.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
GREEN = RGBColor(22, 163, 74)


def load_font(size=32, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=32, fill="#111827", bold=False, anchor=None, align="left"):
    draw.text(xy, value, font=load_font(size, bold), fill=fill, anchor=anchor, align=align)


def wrap_lines(value, max_chars):
    words = value.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def phone_canvas(title, subtitle=None, bg="#f8fafc"):
    img = Image.new("RGB", (900, 1200), "#ffffff")
    draw = ImageDraw.Draw(img)
    rounded(draw, (210, 35, 690, 1165), 48, "#0f172a")
    rounded(draw, (230, 70, 670, 1130), 34, bg)
    rounded(draw, (380, 86, 520, 104), 8, "#111827")
    text(draw, (450, 145), title, 42, "#0f172a", True, anchor="mm")
    if subtitle:
        text(draw, (450, 195), subtitle, 24, "#64748b", False, anchor="mm")
    return img, draw


def save_login(path):
    img, draw = phone_canvas("RoundTracker", "Cada ronda cuenta", "#eef7f0")
    rounded(draw, (275, 275, 625, 345), 16, "#ffffff", "#dbe4ef", 3)
    text(draw, (305, 300), "Usuario", 22, "#64748b")
    text(draw, (305, 323), "iona", 28, "#111827", True)
    rounded(draw, (275, 375, 625, 445), 16, "#ffffff", "#dbe4ef", 3)
    text(draw, (305, 400), "Contraseña", 22, "#64748b")
    text(draw, (305, 423), "••••••••", 28, "#111827", True)
    rounded(draw, (275, 495, 625, 565), 18, "#16a34a")
    text(draw, (450, 530), "ENTRAR", 28, "#ffffff", True, anchor="mm")
    rounded(draw, (275, 645, 625, 790), 18, "#ffffff", "#bbf7d0", 2)
    text(draw, (450, 680), "¿Olvidaste la contraseña?", 23, "#166534", True, anchor="mm")
    text(draw, (450, 725), "Pide un reset al admin", 22, "#64748b", anchor="mm")
    text(draw, (450, 760), "o usa recuperación verificada", 22, "#64748b", anchor="mm")
    img.save(path)


def save_calendar(path):
    img, draw = phone_canvas("Mis torneos", "Calendario personal", "#f8fafc")
    cards = [
        ("Campeonato de España", "La Manga Campo Sur", "#dbeafe"),
        ("Golf Sant Cugat", "Tarjeta roja validada", "#fee2e2"),
        ("Catalunya Infantil", "Clasificación compartida", "#dcfce7"),
    ]
    y = 260
    for title_v, sub, fill in cards:
        rounded(draw, (265, y, 635, y + 130), 20, "#ffffff", "#e2e8f0", 2)
        rounded(draw, (285, y + 20, 335, y + 70), 14, fill)
        text(draw, (355, y + 32), title_v, 26, "#111827", True)
        text(draw, (355, y + 70), sub, 21, "#64748b")
        text(draw, (595, y + 65), "›", 42, "#94a3b8", True, anchor="mm")
        y += 165
    rounded(draw, (300, 825, 600, 895), 18, "#2563eb")
    text(draw, (450, 860), "BUSCAR TORNEO", 25, "#ffffff", True, anchor="mm")
    img.save(path)


def save_join(path):
    img, draw = phone_canvas("Añadirse", "Buscar y apuntarse", "#ffffff")
    rounded(draw, (265, 250, 635, 325), 18, "#f8fafc", "#cbd5e1", 3)
    text(draw, (295, 292), "Buscar torneo...", 25, "#94a3b8")
    rounded(draw, (265, 375, 635, 610), 22, "#f1f5f9", "#dbeafe", 3)
    text(draw, (295, 420), "Campeonato de España", 27, "#111827", True)
    for i, line in enumerate(wrap_lines("Real La Manga Club / La Serena Golf", 28)):
        text(draw, (295, 470 + i * 32), line, 22, "#64748b")
    text(draw, (295, 545), "26-28 junio 2026", 22, "#2563eb", True)
    rounded(draw, (385, 645, 635, 712), 16, "#16a34a")
    text(draw, (510, 678), "APUNTARSE", 24, "#ffffff", True, anchor="mm")
    text(draw, (450, 805), "Al apuntarte apareces\nen la clasificación conjunta.", 27, "#334155", True, anchor="mm", align="center")
    img.save(path)


def save_scorecard(path):
    img, draw = phone_canvas("Resultados", "Ronda 1", "#f8fafc")
    text(draw, (275, 245), "Tarjeta", 30, "#111827", True)
    holes = list(range(1, 10))
    x0, y0 = 270, 310
    cell = 40
    for idx, h in enumerate(holes):
        x = x0 + idx * cell
        text(draw, (x + 20, y0), str(h), 20, "#334155", True, anchor="mm")
        text(draw, (x + 20, y0 + 40), str([4, 3, 5, 4, 4, 3, 5, 4, 4][idx]), 22, "#64748b", True, anchor="mm")
        rounded(draw, (x, y0 + 70, x + 38, y0 + 110), 2, "#fffaf0", "#d6d3d1", 1)
        text(draw, (x + 19, y0 + 91), "-", 22, "#94a3b8", anchor="mm")
    rounded(draw, (275, 520, 625, 590), 18, "#0f172a")
    text(draw, (450, 555), "MODO MÓVIL", 26, "#ffffff", True, anchor="mm")
    rounded(draw, (275, 625, 625, 695), 18, "#16a34a")
    text(draw, (450, 660), "GUARDAR RESULTADOS", 24, "#ffffff", True, anchor="mm")
    rounded(draw, (275, 760, 625, 880), 18, "#ecfdf5", "#bbf7d0", 3)
    text(draw, (450, 800), "Consejo", 24, "#166534", True, anchor="mm")
    text(draw, (450, 842), "Guarda al terminar cada ronda.", 22, "#334155", anchor="mm")
    img.save(path)


def save_mobile(path):
    img, draw = phone_canvas("Modo móvil", "Marcar desde el campo", "#eafaf1")
    text(draw, (450, 265), "Hoyo 7", 52, "#111827", True, anchor="mm")
    text(draw, (450, 320), "Par 5", 30, "#64748b", True, anchor="mm")
    rounded(draw, (300, 390, 600, 610), 40, "#ffffff", "#bbf7d0", 4)
    text(draw, (450, 465), "Golpes", 30, "#64748b", True, anchor="mm")
    text(draw, (450, 545), "5", 70, "#16a34a", True, anchor="mm")
    rounded(draw, (275, 700, 395, 775), 18, "#ffffff", "#cbd5e1", 3)
    text(draw, (335, 737), "‹", 45, "#334155", True, anchor="mm")
    rounded(draw, (505, 700, 625, 775), 18, "#2563eb")
    text(draw, (565, 737), "›", 45, "#ffffff", True, anchor="mm")
    text(draw, (450, 885), "Si se cierra la app,\nvuelve al torneo y continúa.", 26, "#334155", anchor="mm", align="center")
    img.save(path)


def save_leaderboard(path):
    img, draw = phone_canvas("Clasificación", "En directo", "#111827")
    text(draw, (450, 230), "🏆 Leaderboard", 34, "#ffffff", True, anchor="mm")
    headers = ["POS", "JUGADORA", "TOTAL", "VS PAR"]
    xs = [285, 375, 535, 610]
    for x, h in zip(xs, headers):
        text(draw, (x, 310), h, 17, "#94a3b8", True, anchor="mm")
    rows = [
        ("1", "Nicole", "72", "E", "#1e293b"),
        ("2", "Ona", "74", "+2", "#172554"),
        ("-", "Ion Lage", "-", "-", "#1e293b"),
    ]
    y = 360
    for pos, name, total, vs, fill in rows:
        rounded(draw, (260, y, 640, y + 95), 12, fill, "#334155", 1)
        text(draw, (285, y + 48), pos, 25, "#cbd5e1", True, anchor="mm")
        text(draw, (375, y + 35), name, 26, "#ffffff", True, anchor="mm")
        if total == "-":
            text(draw, (375, y + 67), "Sin resultados", 18, "#94a3b8", anchor="mm")
        text(draw, (535, y + 48), total, 26, "#ffffff", True, anchor="mm")
        text(draw, (610, y + 48), vs, 26, "#ffffff", True, anchor="mm")
        y += 115
    rounded(draw, (285, 835, 615, 910), 18, "#2563eb")
    text(draw, (450, 872), "COMPARTIR EN VIVO", 23, "#ffffff", True, anchor="mm")
    img.save(path)


def save_share(path):
    img, draw = phone_canvas("Compartir", "WhatsApp o enlace", "#f8fafc")
    rounded(draw, (275, 290, 625, 420), 24, "#dcfce7", "#86efac", 3)
    text(draw, (450, 335), "Compartir resultados", 27, "#166534", True, anchor="mm")
    text(draw, (450, 375), "Link público en directo", 22, "#334155", anchor="mm")
    rounded(draw, (300, 520, 600, 620), 22, "#25d366")
    text(draw, (450, 570), "WhatsApp", 32, "#ffffff", True, anchor="mm")
    rounded(draw, (300, 690, 600, 790), 22, "#ffffff", "#cbd5e1", 3)
    text(draw, (450, 740), "Copiar enlace", 28, "#334155", True, anchor="mm")
    img.save(path)


def save_profile(path):
    img, draw = phone_canvas("Perfil", "Datos de jugadora", "#ffffff")
    rounded(draw, (370, 235, 530, 395), 80, "#dbeafe", "#60a5fa", 5)
    text(draw, (450, 315), "IR", 54, "#1d4ed8", True, anchor="mm")
    labels = [("Nombre", "Ion Lage"), ("Licencia", "123456"), ("Hándicap", "14.5"), ("Móvil", "+34 *** *** 222")]
    y = 455
    for label, value in labels:
        text(draw, (285, y), label, 20, "#64748b")
        text(draw, (285, y + 33), value, 27, "#111827", True)
        y += 95
    rounded(draw, (275, 875, 625, 945), 18, "#2563eb")
    text(draw, (450, 910), "GUARDAR CAMBIOS", 23, "#ffffff", True, anchor="mm")
    img.save(path)


def save_recovery(path):
    img, draw = phone_canvas("Recuperación", "Dos métodos", "#f8fafc")
    rounded(draw, (275, 260, 625, 395), 20, "#ffffff", "#dbeafe", 3)
    text(draw, (315, 305), "1", 34, "#2563eb", True)
    text(draw, (360, 302), "Código por SMS", 27, "#111827", True)
    text(draw, (360, 340), "Móvil verificado", 21, "#64748b")
    rounded(draw, (275, 450, 625, 585), 20, "#ffffff", "#dbeafe", 3)
    text(draw, (315, 495), "2", 34, "#2563eb", True)
    text(draw, (360, 492), "Email real", 27, "#111827", True)
    text(draw, (360, 530), "Padres o jugadora", 21, "#64748b")
    rounded(draw, (275, 665, 625, 815), 20, "#fff7ed", "#fed7aa", 3)
    text(draw, (450, 705), "Sin móvil/email", 26, "#9a3412", True, anchor="mm")
    text(draw, (450, 755), "El admin genera una\ncontraseña temporal.", 22, "#334155", anchor="mm", align="center")
    img.save(path)


def set_run_font(run, size=None, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "RoundTracker · Manual de uso"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Uso para jugadoras, familias y administradores"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Manual de uso\nRoundTracker")
    set_run_font(r, 30, INK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Cómo registrar torneos, resultados y seguir la clasificación en directo")
    set_run_font(r, 14, MUTED)

    doc.add_picture(str(IMG_DIR / "leaderboard.png"), width=Inches(3.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Versión inicial · Junio 2026")
    set_run_font(r, 10, MUTED, True)
    doc.add_page_break()


def add_step_page(doc, number, title_v, purpose, img_name, steps, note=None):
    h = doc.add_paragraph(style="Heading 1")
    h.add_run(f"{number}. {title_v}")

    p = doc.add_paragraph()
    r = p.add_run(purpose)
    set_run_font(r, 11, INK, True)

    doc.add_picture(str(IMG_DIR / img_name), width=Inches(2.55))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Pasos", style="Heading 2")
    for step in steps:
        doc.add_paragraph(step, style="List Bullet")

    if note:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F4F6F9")
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run("Nota: ")
        set_run_font(r, 10.5, DARK_BLUE, True)
        r = p.add_run(note)
        set_run_font(r, 10.5, INK)

    doc.add_page_break()


def add_faq(doc):
    doc.add_paragraph("10. Preguntas frecuentes", style="Heading 1")
    faqs = [
        ("No veo mi torneo", "Busca el torneo desde la comunidad o pide al administrador que revise si estás apuntada."),
        ("No aparezco en la clasificación", "Apareces cuando estás dada de alta en el torneo. Si aún no has guardado golpes, verás “Sin resultados”."),
        ("El campo o la tarjeta no coincide", "Avísalo antes de introducir toda la ronda. Algunos clubs tienen circuitos distintos según fecha o torneo."),
        ("He olvidado la contraseña", "Pide al administrador una contraseña temporal o usa recuperación por móvil/email cuando esté activada."),
        ("Se cerró la app en el campo", "Vuelve a abrir el torneo y continúa desde la ronda. Conviene guardar al terminar la ronda."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run(q)
        set_run_font(r, 12, DARK_BLUE, True)
        p = doc.add_paragraph(a)
        p.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph("Checklist rápido", style="Heading 2")
    for item in [
        "Entrar con usuario corto, no con email.",
        "Confirmar que el torneo aparece en “Mis torneos”.",
        "Revisar campo y tarjeta antes de guardar resultados.",
        "Usar “Modo móvil” durante el recorrido si se quiere marcar hoyo a hoyo.",
        "Compartir el enlace live cuando la clasificación esté lista.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def build():
    OUT_DIR.mkdir(exist_ok=True)
    IMG_DIR.mkdir(exist_ok=True)

    image_builders = {
        "login.png": save_login,
        "calendar.png": save_calendar,
        "join.png": save_join,
        "scorecard.png": save_scorecard,
        "mobile.png": save_mobile,
        "leaderboard.png": save_leaderboard,
        "share.png": save_share,
        "profile.png": save_profile,
        "recovery.png": save_recovery,
    }
    for filename, builder in image_builders.items():
        builder(IMG_DIR / filename)

    doc = Document()
    style_doc(doc)
    add_title_page(doc)

    add_step_page(
        doc,
        1,
        "Inicio de sesión",
        "Entra con el usuario corto asignado por el equipo.",
        "login.png",
        [
            "Escribe tu usuario sin espacios. Ejemplo: iona.",
            "Introduce la contraseña asignada.",
            "Pulsa Entrar.",
            "Si no recuerdas la contraseña, pide reset al administrador.",
        ],
        "El email interno tipo iona@golfteam.app se usa solo para Firebase. La jugadora no necesita escribirlo.",
    )
    add_step_page(
        doc,
        2,
        "Pantalla principal",
        "Desde aquí se revisan los torneos y se abre la tarjeta de cada ronda.",
        "calendar.png",
        [
            "Revisa la lista de torneos próximos.",
            "Pulsa un torneo para abrir el detalle.",
            "Comprueba fecha, campo y rondas antes de introducir resultados.",
            "Usa Buscar torneo si todavía no aparece en tu calendario.",
        ],
    )
    add_step_page(
        doc,
        3,
        "Añadirse a un torneo",
        "Los torneos compartidos permiten clasificación conjunta entre jugadoras.",
        "join.png",
        [
            "Abre el buscador o la comunidad de torneos.",
            "Busca por nombre, fecha o club.",
            "Pulsa Apuntarse.",
            "Vuelve a Mis torneos y confirma que aparece.",
        ],
        "Si hay dos torneos parecidos, usa el que tenga el campo correcto y el roster común.",
    )
    add_step_page(
        doc,
        4,
        "Introducir resultados",
        "La tarjeta guarda golpes por ronda y alimenta la clasificación en directo.",
        "scorecard.png",
        [
            "Abre el torneo y elige la ronda.",
            "Introduce los golpes de cada hoyo.",
            "Guarda los resultados al terminar.",
            "Revisa que la suma y el par tengan sentido.",
        ],
        "Antes de guardar muchas rondas, valida que la tarjeta del campo corresponde al circuito del torneo.",
    )
    add_step_page(
        doc,
        5,
        "Modo móvil en campo",
        "Vista simplificada para marcar golpes durante el recorrido.",
        "mobile.png",
        [
            "Pulsa Modo móvil desde la tarjeta.",
            "Marca los golpes del hoyo actual.",
            "Avanza o retrocede con los controles.",
            "Al terminar, vuelve a la tarjeta y guarda.",
        ],
    )
    add_step_page(
        doc,
        6,
        "Clasificación en directo",
        "La clasificación muestra resultados guardados y jugadoras pendientes.",
        "leaderboard.png",
        [
            "Abre la pestaña Clasificación.",
            "Revisa posición, total y vs par.",
            "Las jugadoras sin tarjeta guardada aparecen como Sin resultados.",
            "Actualiza la vista si acabas de guardar y no se refleja todavía.",
        ],
    )
    add_step_page(
        doc,
        7,
        "Compartir resultados",
        "El enlace live permite que familiares o entrenadores vean la clasificación.",
        "share.png",
        [
            "Pulsa Compartir o En vivo.",
            "Elige WhatsApp o copiar enlace.",
            "Envía el enlace al grupo o a la familia.",
            "Quien abre el enlace no necesita contraseña.",
        ],
        "Comparte solo enlaces públicos de torneo. La edición de resultados sigue protegida por login.",
    )
    add_step_page(
        doc,
        8,
        "Perfil y datos",
        "Mantén actualizados nombre, licencia, hándicap y datos de recuperación.",
        "profile.png",
        [
            "Abre el perfil desde la app.",
            "Comprueba nombre y licencia.",
            "Actualiza hándicap si corresponde.",
            "Añade móvil o email real cuando el sistema de recuperación esté activo.",
        ],
    )
    add_step_page(
        doc,
        9,
        "Recuperar contraseña",
        "La recuperación debe usar móvil o email real, no el email interno de golfteam.app.",
        "recovery.png",
        [
            "Si tienes móvil verificado, pide código SMS.",
            "Si tienes email real verificado, pide código o enlace por email.",
            "Si no tienes método configurado, pide reset al administrador.",
            "Cuando recibas una contraseña temporal, cámbiala al entrar.",
        ],
        "Las contraseñas antiguas no se pueden consultar. Solo se pueden cambiar por una nueva.",
    )
    add_faq(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
